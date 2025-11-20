#!/usr/bin/env python3
"""
Extrai e Organiza TODOS os recursos dos ZIPs
Converte tudo para PDF e cria índice completo para os agentes IA
"""

import os
import json
import shutil
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet

def extract_text_from_docx(docx_path):
    """Extrai texto de um arquivo DOCX"""
    try:
        doc = Document(docx_path)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extrair tabelas
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"   ⚠️ Erro ao ler DOCX: {e}")
        return ""

def convert_docx_to_pdf(docx_path, pdf_path):
    """Converte DOCX para PDF usando ReportLab"""
    try:
        # Extrair texto do DOCX
        text = extract_text_from_docx(docx_path)
        
        if not text:
            print(f"   ⚠️ Sem texto para converter")
            return False
        
        # Criar PDF
        doc = SimpleDocTemplate(str(pdf_path), pagesize=letter, 
                               topMargin=0.75*inch, bottomMargin=0.75*inch)
        story = []
        styles = getSampleStyleSheet()
        
        # Título (nome do arquivo)
        title = Path(docx_path).stem.replace('_', ' ').title()
        story.append(Paragraph(f"<b>{title}</b>", styles['Heading1']))
        story.append(Spacer(1, 0.3*inch))
        
        # Conteúdo
        paragraphs = text.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                # Detectar se é título
                if len(para_text) < 100 and para_text.isupper():
                    story.append(Paragraph(f"<b>{para_text}</b>", styles['Heading2']))
                else:
                    story.append(Paragraph(para_text, styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        doc.build(story)
        return True
        
    except Exception as e:
        print(f"   ❌ Erro na conversão: {e}")
        return False

def organize_all_resources():
    """Organiza todos os recursos dos ZIPs"""
    
    print("\n" + "="*80)
    print("📦 EXTRAINDO E ORGANIZANDO TODOS OS RECURSOS")
    print("="*80)
    
    # Diretórios
    base_dir = Path("/app/immigration_resources")
    base_dir.mkdir(exist_ok=True)
    
    # Categorias
    categories = {
        "forms": base_dir / "forms_and_templates",
        "guides": base_dir / "completion_guides",
        "checklists": base_dir / "checklists_and_trackers",
        "templates": base_dir / "letter_templates",
        "official": base_dir / "official_forms_pdf"
    }
    
    for cat_dir in categories.values():
        cat_dir.mkdir(exist_ok=True)
    
    # Processar arquivos do diretório oficial_forms
    source_dir = Path("/app/official_forms")
    
    if not source_dir.exists():
        print("❌ Diretório oficial_forms não encontrado")
        return
    
    all_files = list(source_dir.glob("**/*.*"))
    print(f"\n📄 Total de arquivos encontrados: {len(all_files)}")
    
    # Catalogar por tipo
    catalog = {
        "forms_and_templates": [],
        "completion_guides": [],
        "checklists_and_trackers": [],
        "letter_templates": [],
        "official_forms": []
    }
    
    converted_count = 0
    copied_count = 0
    
    for file_path in all_files:
        if not file_path.is_file():
            continue
        
        filename = file_path.name
        
        # Determinar categoria
        category = None
        if "Guide" in filename or "Completion" in filename:
            category = "guides"
            catalog_key = "completion_guides"
        elif "Checklist" in filename or "Tracker" in filename or "Log" in filename:
            category = "checklists"
            catalog_key = "checklists_and_trackers"
        elif "Letter" in filename or "Affidavit" in filename:
            category = "templates"
            catalog_key = "letter_templates"
        elif "Form" in filename or "Index" in filename or "Flowchart" in filename:
            category = "forms"
            catalog_key = "forms_and_templates"
        else:
            category = "forms"
            catalog_key = "forms_and_templates"
        
        # Processar arquivo
        dest_dir = categories[category]
        
        if file_path.suffix.lower() == '.docx':
            # Converter DOCX para PDF
            pdf_name = file_path.stem + '.pdf'
            pdf_path = dest_dir / pdf_name
            
            print(f"\n🔄 Convertendo: {filename}")
            if convert_docx_to_pdf(file_path, pdf_path):
                print(f"   ✅ PDF criado: {pdf_name}")
                catalog[catalog_key].append({
                    "filename": pdf_name,
                    "original": filename,
                    "path": str(pdf_path.relative_to(base_dir)),
                    "type": "PDF (convertido de DOCX)"
                })
                converted_count += 1
            
        elif file_path.suffix.lower() == '.pdf':
            # Copiar PDF
            dest_path = categories["official"] / filename
            shutil.copy2(file_path, dest_path)
            catalog["official_forms"].append({
                "filename": filename,
                "path": str(dest_path.relative_to(base_dir)),
                "type": "PDF (oficial)"
            })
            copied_count += 1
            print(f"   ✅ Copiado: {filename}")
    
    # Salvar catálogo
    catalog_file = base_dir / "catalog.json"
    with open(catalog_file, 'w', encoding='utf-8') as f:
        json.dump(catalog, f, indent=2, ensure_ascii=False)
    
    print(f"\n" + "="*80)
    print(f"✅ ORGANIZAÇÃO CONCLUÍDA!")
    print(f"="*80)
    print(f"📊 Estatísticas:")
    print(f"   • DOCXs convertidos para PDF: {converted_count}")
    print(f"   • PDFs copiados: {copied_count}")
    print(f"   • Total de recursos: {converted_count + copied_count}")
    print(f"\n📁 Estrutura criada:")
    print(f"   • Forms & Templates: {len(catalog['forms_and_templates'])} arquivos")
    print(f"   • Completion Guides: {len(catalog['completion_guides'])} arquivos")
    print(f"   • Checklists & Trackers: {len(catalog['checklists_and_trackers'])} arquivos")
    print(f"   • Letter Templates: {len(catalog['letter_templates'])} arquivos")
    print(f"   • Official Forms: {len(catalog['official_forms'])} arquivos")
    print(f"\n📖 Catálogo salvo: {catalog_file}")
    print(f"="*80)
    
    return catalog

def create_master_index():
    """Cria índice master em PDF para fácil consulta"""
    
    base_dir = Path("/app/immigration_resources")
    catalog_file = base_dir / "catalog.json"
    
    if not catalog_file.exists():
        print("❌ Catálogo não encontrado")
        return
    
    with open(catalog_file, 'r', encoding='utf-8') as f:
        catalog = json.load(f)
    
    # Criar PDF do índice
    index_pdf = base_dir / "MASTER_INDEX.pdf"
    doc = SimpleDocTemplate(str(index_pdf), pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Título
    story.append(Paragraph("<b>IMMIGRATION RESOURCES MASTER INDEX</b>", styles['Title']))
    story.append(Spacer(1, 0.5*inch))
    
    story.append(Paragraph(
        "This index contains all immigration resources available for AI agents to consult, "
        "including forms, templates, guides, checklists, and official USCIS documents.",
        styles['Normal']
    ))
    story.append(Spacer(1, 0.3*inch))
    
    # Seções
    sections = {
        "completion_guides": "Completion Guides & Instructions",
        "checklists_and_trackers": "Checklists & Case Trackers",
        "letter_templates": "Letter & Document Templates",
        "forms_and_templates": "Forms & Master Templates",
        "official_forms": "Official USCIS Forms (PDF)"
    }
    
    for key, title in sections.items():
        if catalog.get(key):
            story.append(Paragraph(f"<b>{title}</b>", styles['Heading1']))
            story.append(Spacer(1, 0.2*inch))
            
            for i, item in enumerate(catalog[key], 1):
                filename = item.get('filename', item.get('original', 'Unknown'))
                path = item.get('path', '')
                
                story.append(Paragraph(
                    f"{i}. <b>{filename}</b>",
                    styles['Normal']
                ))
                story.append(Paragraph(
                    f"   Path: {path}",
                    styles['Normal']
                ))
                story.append(Spacer(1, 0.15*inch))
            
            story.append(PageBreak())
    
    doc.build(story)
    
    print(f"\n✅ Índice master criado: {index_pdf}")
    print(f"   Total de seções: {len(sections)}")
    print(f"   Total de recursos indexados: {sum(len(catalog.get(k, [])) for k in sections.keys())}")
    
    return index_pdf

if __name__ == "__main__":
    # Organizar todos os recursos
    catalog = organize_all_resources()
    
    # Criar índice master
    if catalog:
        index_file = create_master_index()
        
        print(f"\n🎉 TUDO PRONTO!")
        print(f"\n📚 Recursos disponíveis em: /app/immigration_resources/")
        print(f"📖 Consulte o MASTER_INDEX.pdf para navegação completa")
